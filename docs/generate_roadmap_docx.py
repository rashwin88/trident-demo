"""Generate docs/trident-roadmap.docx from the full Trident roadmap.

Covers all three axes:
    - Representation (R1-R9)
    - Querying     (Q1-Q10)
    - SLMs         (S1-S5)

Run: python3 docs/generate_roadmap_docx.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()


# ── Helpers ──────────────────────────────────


def add_section(num, title, type_, effort, impact):
    doc.add_page_break()
    doc.add_heading(f"{num}. {title}", level=2)
    meta = doc.add_paragraph()
    meta.add_run("Type: ").bold = True
    meta.add_run(f"{type_}   ")
    meta.add_run("Effort: ").bold = True
    meta.add_run(f"{effort}   ")
    meta.add_run("Impact: ").bold = True
    meta.add_run(impact)


def add_subheading(text):
    doc.add_heading(text, level=3)


def add_para(text):
    doc.add_paragraph(text)


def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_axis_heading(text):
    doc.add_page_break()
    h = doc.add_heading(text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_matrix(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v


# ── Cover page ───────────────────────────────

title = doc.add_heading("Trident — Technical Roadmap", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Representation  ·  Querying  ·  Distilled SLMs")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()
doc.add_heading("Overview", level=1)
doc.add_paragraph(
    "This document captures the complete Trident technical roadmap across three axes:"
)
add_bullets([
    "Axis 1 — Representation (R1-R9): how knowledge is stored and structured in the graph and vector layers.",
    "Axis 2 — Querying (Q1-Q10): how the LangGraph agent finds and reasons over knowledge, including new tools and query patterns.",
    "Axis 3 — Distilled SLMs (S1-S5): task-specific small language models that replace frontier calls for extraction, classification, evaluation, Cypher generation, and entity deduplication.",
])
doc.add_paragraph(
    "Each item is tagged with effort (Low / Medium / High) and impact. The document is structured so each "
    "axis can be consumed independently — see the Priority Matrix at the top of each axis for a quick "
    "summary before drilling into the detail."
)


# ──────────────────────────────────────────────
# AXIS 1 — REPRESENTATION
# ──────────────────────────────────────────────

add_axis_heading("Axis 1 — Representation (Knowledge Structure)")

doc.add_paragraph(
    "Representation changes are the scaffolding that everything else depends on. The priority order favours "
    "fixing correctness bugs first (R1), then adding signal that downstream querying can exploit (R2-R6), "
    "then scaling up to global/aggregate structures (R4, R9)."
)

doc.add_heading("Priority Matrix", level=2)
add_matrix(
    ["#", "Improvement", "Effort", "Impact", "Dependencies", "Status"],
    [
        ("R1", "Fix chunk-scoped entity edges", "Low", "High", "None (bug fix)", "Completed"),
        ("R2", "Graph-boosted chunk reranking", "Low", "High", "Best after R1", "Pending"),
        ("R3", "Temporal metadata", "Low", "Medium", "None", "Pending"),
        ("R4", "Community detection + summaries", "Medium", "High", "None", "Pending"),
        ("R5", "Proposition confidence scoring", "Low", "Medium", "None", "Pending"),
        ("R6", "Edge weights + time decay", "Low", "Medium", "None", "Pending"),
        ("R7", "Adaptive entity resolution", "Medium", "Medium", "None", "Pending"),
        ("R8", "Embedding drift detection", "Medium", "Medium", "None", "Pending"),
        ("R9", "DRIFT-style community search", "Low", "Medium", "Requires R4", "Pending"),
    ],
)

# ── R1 ───────────────────────────────────────
add_section("R1", "Fix Chunk-Scoped Entity Edges", "Bug fix", "Low", "High — fixes noisy graph traversals")
add_subheading("Problem")
add_para(
    "Stage 5d-5e in the pipeline created MENTIONS edges from every chunk to every entity and DEFINES edges "
    "from every chunk to every concept — regardless of whether that entity/concept was actually extracted "
    "from that chunk. This was a Cartesian product. With 20 chunks and 50 entities: 1,000 edges instead of "
    "the ~100 that are valid. This drowned meaningful relationships in Chunk noise, made BFS expansion "
    "return irrelevant nodes, broke graph-boosted reranking, and inflated graph statistics."
)
add_subheading("Fix")
add_para(
    "Thread chunk_id through the extraction result so each entity/concept knows which chunk it came from. "
    "In the Store stage, only create MENTIONS edges between a chunk and entities extracted from that chunk, "
    "and DEFINES edges between a chunk and concepts extracted from that chunk."
)
add_subheading("Files Changed")
add_bullets([
    "backend/ingestion/extractor.py — attach chunk_id to each extracted entity/concept",
    "backend/ingestion/pipeline.py — Stage 5d/5e: use per-chunk entity lists instead of all_entities",
])

# ── R2 ───────────────────────────────────────
add_section("R2", "Graph-Boosted Chunk Reranking", "Enhancement", "Low (~10 lines)", "High — dramatically improves retrieval precision")
add_subheading("Problem")
add_para(
    "The query engine runs GN search (graph anchoring) and KS search (chunk retrieval) independently. "
    "Results are concatenated with no cross-signal. A chunk that mentions 3 anchor entities ranks the "
    "same as one that mentions none."
)
add_subheading("Fix")
add_para("After Step 4 (chunk retrieval) in query/engine.py, add a reranking step:")
code = doc.add_paragraph()
code.add_run(
    "for chunk in ks_results:\n"
    "    graph_overlap = await graph.count_chunk_anchor_overlap(\n"
    "        chunk.chunk_id, anchor_ids, provider_id\n"
    "    )\n"
    "    chunk.final_score = 0.6 * chunk.vector_score + 0.4 * (graph_overlap / max(len(anchor_ids), 1))\n\n"
    "ks_results.sort(key=lambda c: c.final_score, reverse=True)"
).font.name = "Menlo"
add_para("Supporting Cypher:")
cypher = doc.add_paragraph()
cypher.add_run(
    "MATCH (c:Chunk {chunk_id: $cid, provider_id: $pid})-[:MENTIONS|DEFINES]->(n)\n"
    "WHERE elementId(n) IN $anchor_ids\n"
    "RETURN count(n) AS overlap"
).font.name = "Menlo"
add_subheading("Files to Change")
add_bullets([
    "backend/stores/graph.py — add count_chunk_anchor_overlap() method",
    "backend/query/engine.py — add reranking between Step 4 and Step 6",
])
add_subheading("References")
add_bullets([
    "GraphER: Graph-Based Enrichment and Reranking for RAG",
    "Graph-Augmented Hybrid Retrieval and Multi-Stage Re-ranking",
])

# ── R3 ───────────────────────────────────────
add_section("R3", "Temporal Metadata (Bi-Temporal Model)", "Enhancement", "Low", "Medium — prevents stale fact accumulation")
add_subheading("Problem")
add_para(
    "Nodes have no timestamps. Re-ingesting an updated document creates duplicate facts. No way to know "
    "what is current vs outdated. Agents can't answer 'what changed?' or 'what was true before?'"
)
add_subheading("Fix")
add_para("Add properties to nodes:")
add_bullets([
    "All nodes: created_at (ingestion timestamp), source_doc_hash (SHA-256 of source file)",
    "Propositions: valid_from (datetime), valid_until (datetime, nullable)",
])
add_para("On re-ingest (same filename, same provider):")
add_bullets([
    "1. Hash the new file content.",
    "2. If hash differs from existing document node's source_doc_hash: mark existing propositions from that document as valid_until = datetime(). New propositions get valid_from = datetime().",
    "3. If hash matches: skip (idempotent).",
])
add_para("Query-time filter: WHERE p.valid_until IS NULL (default). Optional include_historical=true.")
add_subheading("Files to Change")
add_bullets([
    "backend/stores/graph.py — add timestamps to merge_entity, create_proposition_node, etc.",
    "backend/ingestion/pipeline.py — add document hash check at parse stage",
    "backend/query/engine.py — add temporal filter to graph context",
])
add_subheading("References")
add_bullets([
    "Zep: A Temporal Knowledge Graph Architecture for Agent Memory",
    "Graphiti: Knowledge Graph Memory for an Agentic World",
])

# ── R4 ───────────────────────────────────────
add_section("R4", "Community Detection + Hierarchical Summaries", "New feature", "Medium", "High — enables global/thematic queries")
add_subheading("Problem")
add_para(
    "The system handles local queries ('what is entity X?') well but fails on global ones ('what are the "
    "main themes?', 'summarize everything about compliance'). No aggregation or clustering exists."
)
add_subheading("Fix — New pipeline stage (post-Store)")
add_bullets([
    "1. Pull the entity/concept/relationship subgraph from Neo4j.",
    "2. Run the Leiden algorithm (via graspologic or python-igraph) at 2-3 resolution levels.",
    "3. Create Community nodes in Neo4j with level, member_count, summary properties.",
    "4. Add MEMBER_OF edges from entities/concepts to their community.",
    "5. LLM-summarize each community (DSPy CommunitySignature) — concatenate member descriptions → summary.",
    "6. Embed community summaries in GN index with node_type='Community'.",
])
add_para(
    "Query-time: When GN search returns Community hits above threshold, include the summary in "
    "graph_context for thematic framing."
)
add_subheading("New Dependencies")
add_bullets(["graspologic>=3.0 or python-igraph>=0.11"])
add_subheading("Cost")
add_para("~3-5 extra LLM calls per ingestion (one per community). Negligible at query time.")
add_subheading("References")
add_bullets([
    "From Local to Global: A Graph RAG Approach to Query-Focused Summarization",
    "Microsoft GraphRAG Community Detection",
])

# ── R5 ───────────────────────────────────────
add_section("R5", "Proposition Confidence Scoring", "Enhancement", "Low", "Medium — enables fact reliability ranking")
add_subheading("Problem")
add_para(
    "A proposition mentioned in 5 documents is treated the same as one mentioned once. The LLM has no "
    "signal for fact reliability."
)
add_subheading("Fix")
add_bullets([
    "Add confidence float (default 1.0) to Proposition nodes.",
    "During Resolve: when a new proposition's SPO embedding matches an existing one (cosine ≥ 0.90), increment confidence by 0.5 instead of creating a duplicate. Add CORROBORATED_BY edge to the new source chunk.",
    "In _format_graph_context: include confidence so the LLM prefers well-corroborated facts.",
    "Embed confidence as a filterable field in GN index.",
])
add_subheading("Files to Change")
add_bullets([
    "backend/models.py — add confidence to proposition model",
    "backend/ingestion/resolver.py — add proposition resolution",
    "backend/stores/graph.py — update create_proposition_node to check for duplicates",
    "backend/query/engine.py — include confidence in context",
])

# ── R6 ───────────────────────────────────────
add_section("R6", "Edge Weights + Time Decay", "Enhancement", "Low", "Medium — improves graph traversal quality over time")
add_subheading("Problem")
add_para(
    "All edges are binary. A relationship mentioned 50 times ranks the same as one mentioned once during "
    "BFS. Stale relationships never get deprioritized."
)
add_subheading("Fix")
add_bullets([
    "Add weight (float, default 1.0) and last_referenced (datetime) to all relationship edges.",
    "On duplicate edge creation: SET r.weight = r.weight + 1, r.last_referenced = datetime().",
    "In BFS/traversal: sort edges by weight, optionally filter WHERE r.weight >= $min_weight.",
    "Periodic decay: SET r.weight = r.weight * 0.95 WHERE r.last_referenced < datetime() - duration('P90D').",
    "Prune edges below 0.1.",
])
add_subheading("Files to Change")
add_bullets([
    "backend/stores/graph.py — update create_edge for weight increment, add apply_decay() method",
    "backend/stores/graph.py — update get_reasoning_subgraph to weight-sort",
    "New: /providers/{id}/maintain endpoint for decay",
])
add_subheading("References")
add_bullets(["Cognee: Knowledge Graph Optimization (Memify)"])

# ── R7 ───────────────────────────────────────
add_section("R7", "Adaptive Entity Resolution (Block + Match)", "Enhancement", "Medium", "Medium — faster and more accurate at scale")
add_subheading("Problem")
add_para(
    "Resolution does N individual embedding + search calls per chunk — slow at scale. Borderline matches "
    "(0.78-0.88) are either all accepted or all rejected by a fixed threshold."
)
add_subheading("Fix — Blocking phase")
add_bullets([
    "Batch-embed all entities from a chunk in one call.",
    "Batch-search GN (Milvus supports this natively) — 1 call instead of N.",
])
add_subheading("Matching phase")
add_bullets([
    "For scores in the borderline range (0.78-0.88), use a lightweight LLM confirmation: 'Are these the same entity? A: {label, desc}. B: {label, desc}. Yes/No.'",
    "Above 0.88: auto-merge. Below 0.78: auto-create.",
])
add_subheading("Audit Trail")
add_bullets([
    "Log merges in MergeEvent nodes with both labels, score, timestamp.",
    "Enables undo of bad merges.",
])
add_subheading("Files to Change")
add_bullets([
    "backend/ingestion/resolver.py — batch embedding, borderline LLM check",
    "backend/stores/graph_index.py — add search_batch() method",
    "backend/stores/graph.py — add MergeEvent node creation",
])

# ── R8 ───────────────────────────────────────
add_section("R8", "Embedding Drift Detection + Re-Indexing", "Operational", "Medium", "Medium — prevents silent degradation on model changes")
add_subheading("Problem")
add_para(
    "Changing the embedding model (e.g., upgrading from text-embedding-3-small) makes all existing "
    "embeddings incompatible with new query embeddings. Silent retrieval quality degradation."
)
add_subheading("Fix")
add_bullets([
    "Store model_name and model_version in each Milvus collection schema.",
    "Add /providers/{id}/reindex endpoint: read all nodes from Neo4j → re-embed with current model → upsert into new collections → swap atomically.",
    "Health check: compare stored model version against configured model. Flag needs_reindex if different.",
    "Log resolution scores during ingestion for threshold calibration over time.",
])

# ── R9 ───────────────────────────────────────
add_section("R9", "DRIFT-Style Community-Aware Search", "Enhancement", "Low (after R4)", "Medium — adds thematic framing to local results")
add_subheading("Problem")
add_para("BFS expansion gives neighbouring nodes but no thematic context.")
add_subheading("Fix")
add_para(
    "Requires Community Detection (R4) first. For each anchor node, traverse up to Community nodes and "
    "include their summaries in context at multiple levels (topic + theme)."
)
add_subheading("Files to Change")
add_bullets(["backend/query/engine.py — add community context step between anchor search and BFS"])
add_subheading("References")
add_bullets([
    "Microsoft DRIFT Search: Dynamic Reasoning and Inference with Flexible Traversal",
    "What is GraphRAG: Complete Guide 2026",
])


# ──────────────────────────────────────────────
# AXIS 2 — QUERYING
# ──────────────────────────────────────────────

add_axis_heading("Axis 2 — Querying (Agent Intelligence)")

doc.add_paragraph(
    "How agents find and reason over knowledge. New tools and query patterns for the LangGraph agent."
)

doc.add_heading("Priority Matrix", level=2)
add_matrix(
    ["#", "Improvement", "Effort", "Impact", "Status"],
    [
        ("Q1", "Query decomposition planner (LangGraph node)", "Medium", "High", "Pending"),
        ("Q2", "trident_find_path — path-based reasoning", "Low", "High", "Pending"),
        ("Q3", "trident_hybrid_search — vector + graph RRF fusion", "Medium", "High", "Pending"),
        ("Q4", "Provenance tracking — source evidence with answers", "Medium", "High", "Pending"),
        ("Q5", "trident_aggregate — pre-built analytics queries", "Low", "Medium", "Pending"),
        ("Q6", "trident_compare — entity side-by-side comparison", "Low", "Medium", "Pending"),
        ("Q7", "Conversation entity memory (Graphiti-inspired)", "Low", "Medium", "Pending"),
        ("Q8", "trident_impact — 'what if' / blast radius analysis", "Medium", "Medium", "Pending"),
        ("Q9", "trident_text2cypher — schema-aware NL→Cypher with retry", "High", "Medium", "Pending"),
        ("Q10", "trident_suggest — graph-aware query suggestions", "Low", "Low", "Pending"),
    ],
)

# ── Q1 ───────────────────────────────────────
add_section("Q1", "Query Decomposition Planner", "LangGraph architecture change", "Medium", "High — improves all multi-hop queries")
add_subheading("Problem")
add_para(
    "Complex questions like 'Which procedures reference entities governed by regulation X?' need to be "
    "broken into sub-queries. The agent currently goes straight from question → tools via ReAct, leading "
    "to aimless exploration on multi-hop questions."
)
add_subheading("Fix")
add_para(
    "Add a planner node before the main agent loop in the LangGraph. The planner receives the question + "
    "schema and outputs a structured plan: ordered sub-questions, each tagged with a suggested tool. The "
    "plan is injected as a system message and becomes visible in the reasoning trace."
)
add_subheading("References")
add_bullets([
    "Generate-on-Graph: LLM as both Agent and KG",
    "Neo4j Multi-Hop Reasoning with Knowledge Graphs and LLMs",
])

# ── Q2 ───────────────────────────────────────
add_section("Q2", "trident_find_path — Path-Based Reasoning", "New tool", "Low", "High — unlocks the core graph differentiator")
add_subheading("Problem")
add_para(
    "'How is Entity A connected to Entity B?' is the killer feature of graph databases that vector search "
    "cannot replicate. Currently requires the agent to write raw Cypher with shortestPath — which LLMs "
    "get wrong frequently."
)
add_subheading("Fix")
add_para(
    "Dedicated tool wrapping Neo4j's allShortestPaths((a)-[*..6]-(b)) with proper provider scoping. "
    "Returns human-readable path strings: A -[RELATED_TO]→ B -[PART_OF]→ C."
)
add_subheading("References")
add_bullets(["Text2Cypher Guide — path syntax is a top LLM failure mode"])

# ── Q3 ───────────────────────────────────────
add_section("Q3", "trident_hybrid_search — Vector + Graph Fused with RRF", "New tool", "Medium", "High — measurable accuracy improvement (~8%)")
add_subheading("Problem")
add_para(
    "Vector search and graph traversal are separate tools. The agent must manually combine results. Nodes "
    "found by both methods are not boosted."
)
add_subheading("Fix")
add_para(
    "Single tool that: (1) runs vector search on GN → (2) takes top-5 hits → (3) expands their graph "
    "neighbourhoods → (4) fuses all results with Reciprocal Rank Fusion: "
    "score = 1/(60 + vector_rank) + 1/(60 + graph_rank). Nodes found by both rank highest."
)
add_subheading("References")
add_bullets([
    "HybridRAG: Integrating Knowledge Graphs and Vector Retrieval",
    "Towards Practical GraphRAG",
])

# ── Q4 ───────────────────────────────────────
add_section("Q4", "Provenance Tracking", "Architecture enhancement", "Medium", "High — trust and explainability")
add_subheading("Problem")
add_para(
    "'The CFO is Ramanathan J' is more useful as 'The CFO is Ramanathan J (source: Entity node via "
    "DESCRIBED_BY edge, corroborated by chunk from leadership.pdf).' Currently no tracking of which "
    "tools/nodes/chunks led to each claim."
)
add_subheading("Fix")
add_para(
    "Extend AgentState with a provenance accumulator. Wrap ToolNode to intercept results and extract "
    "source node_ids, chunk_ids, Cypher queries. Append provenance summary to system prompt before final "
    "answer. Final SSE event includes a sources list."
)
add_subheading("References")
add_bullets(["Paths-over-Graph: KG-Empowered LLM Reasoning"])

# ── Q5 ───────────────────────────────────────
add_section("Q5", "trident_aggregate — Pre-Built Analytics", "New tool", "Low", "Medium — eliminates Cypher errors for analytics")
add_subheading("Problem")
add_para(
    "'How many entities of type X?' or 'Which nodes have the most connections?' require aggregation "
    "Cypher, which LLMs write incorrectly (wrong GROUP BY, missing COLLECT)."
)
add_subheading("Fix")
add_para(
    "Parameterized tool with pre-validated Cypher templates: node_counts, edge_counts, degree_centrality, "
    "entity_cooccurrence, group_by_property. No LLM-generated Cypher — deterministic and fast."
)
add_subheading("References")
add_bullets(["Production-Ready Graph Systems in 2025"])

# ── Q6 ───────────────────────────────────────
add_section("Q6", "trident_compare — Entity Side-by-Side", "New tool", "Low", "Medium — common query pattern, easy win")
add_subheading("Problem")
add_para(
    "'How does Entity A differ from Entity B?' currently requires 4-6 manual tool calls. The agent often "
    "collects asymmetric information."
)
add_subheading("Fix")
add_para(
    "Single tool that extracts both entities' subgraphs in parallel and returns: shared connections, "
    "unique connections per entity, property differences. Two find_exact + get_node_detail calls, then "
    "set intersection/difference."
)

# ── Q7 ───────────────────────────────────────
add_section("Q7", "Conversation Entity Memory", "Enhancement", "Low", "Medium — better multi-turn conversations")
add_subheading("Problem")
add_para(
    "The agent re-searches for entities discussed 5 turns ago. Flat message history loses entity context."
)
add_subheading("Fix")
add_para(
    "Track entities mentioned across turns with recency-decayed scores: entity_score *= 0.8 (decay) + "
    "1.0 (boost on mention). Inject top-5 'focus entities' into system prompt each turn. No external "
    "storage needed."
)
add_subheading("References")
add_bullets([
    "Zep: Temporal Knowledge Graph Architecture for Agent Memory",
    "Graphiti: Knowledge Graph Memory for an Agentic World",
])

# ── Q8 ───────────────────────────────────────
add_section("Q8", "trident_impact — Blast Radius Analysis", "New tool", "Medium", "Medium — high value for operational KBs")
add_subheading("Problem")
add_para(
    "'What would be affected if we removed Entity X?' or 'What depends on this procedure?' — maps to "
    "graph dependency analysis."
)
add_subheading("Fix")
add_para(
    "Tool that finds all nodes reachable from a target within N hops, grouped by hop distance and node "
    "type. Uses APOC subgraphNodes + shortestPath length. Returns "
    "{distance_1: [{type, nodes}], distance_2: [...]}."
)

# ── Q9 ───────────────────────────────────────
add_section("Q9", "trident_text2cypher — Schema-Aware NL→Cypher", "New tool", "High", "Medium — fallback for complex edge cases")
add_subheading("Problem")
add_para(
    "Raw trident_cypher requires perfect Cypher from the LLM. Fails on relationship directions, property "
    "names, syntax."
)
add_subheading("Fix")
add_para(
    "Takes natural language → injects live schema + few-shot examples → generates Cypher → validates with "
    "regex → auto-corrects direction errors → executes → retries once on failure with error context. "
    "Stores validated pairs as few-shot examples that improve over time."
)
add_subheading("References")
add_bullets([
    "Neo4j Text2Cypher Guide",
    "Neo4j Text2Cypher Datasets",
])

# ── Q10 ──────────────────────────────────────
add_section("Q10", "trident_suggest — Query Suggestions", "New tool", "Low", "Low — nice UX, not critical")
add_subheading("Problem")
add_para("Users don't know what's in the graph. 'What can I ask?' returns generic help.")
add_subheading("Fix")
add_para(
    "Tool that surfaces: highest-degree entities (most connected), available relationship patterns, "
    "procedures, entity type distribution. Converts to natural language suggestions like 'You could ask "
    "about [entity] which has [N] connections.'"
)


# ──────────────────────────────────────────────
# AXIS 3 — DISTILLED SLMs
# ──────────────────────────────────────────────

add_axis_heading("Axis 3 — Distilled Small Language Models (SLMs)")

doc.add_paragraph(
    "Trident currently relies on large frontier models (GPT-5, Claude Opus) for every extraction and "
    "reasoning step. Each call costs 2-5K tokens and takes 20-40s. Distilling task-specific SLMs "
    "(3B-8B parameter models fine-tuned on domain-specific traces) can reduce cost by 10-50x and latency "
    "by 5-10x while maintaining quality on narrow, well-scoped tasks."
)
doc.add_paragraph(
    "The strategy: use the frontier model as a teacher to generate high-quality training traces, then "
    "fine-tune a small open-weight model (Phi-3, Llama-3.1-8B, Qwen2.5-7B) on those traces. The SLM "
    "replaces the frontier model for the specific task once it reaches quality parity on the evaluation set."
)

doc.add_heading("Priority Matrix", level=2)
add_matrix(
    ["#", "SLM Target", "Effort", "Impact", "Base Model", "Status"],
    [
        ("S1", "Entity/Relation extractor", "High", "Very High", "Qwen2.5-7B / Llama-3.1-8B", "Pending"),
        ("S2", "Edge-type classifier", "Medium", "High", "Phi-3-mini (3.8B)", "Pending"),
        ("S3", "Extraction quality evaluator", "Medium", "High", "Llama-3.1-8B", "Pending"),
        ("S4", "Cypher query generator", "High", "Medium", "Qwen2.5-Coder-7B", "Pending"),
        ("S5", "Semantic resolver / dedup judge", "Low", "Medium", "Phi-3-mini (3.8B)", "Pending"),
    ],
)

# ── S1 ───────────────────────────────────────
add_section("S1", "Entity/Relation Extraction SLM",
            "Fine-tuned extractor (replaces UnifiedExtractionSignature)",
            "High", "Very High — extraction dominates ingestion cost and latency")
add_subheading("Problem")
add_para(
    "Every document chunk triggers a frontier-model call via DSPy ChainOfThought to extract entities, "
    "concepts, relationships, and propositions. On a 2048-token chunk this is ~4-6K total tokens and "
    "takes 25-40s per chunk. For a 50-page document with 25 chunks, extraction alone costs ~$0.50-$1.50 "
    "and takes 10-15 minutes."
)
add_subheading("Target SLM")
add_para(
    "A 7-8B parameter model (Qwen2.5-7B-Instruct or Llama-3.1-8B-Instruct) fine-tuned to emit the same "
    "Pydantic-typed ExtractionOutput schema. Post-training, runs locally on a single RTX 4090 (24GB) or "
    "A10G cloud instance at ~500-1000 tokens/sec."
)
add_subheading("Training Data")
add_bullets([
    "Source: Frontier-model extraction traces captured during production ingestion — every extract_unified() call's input chunk + raw LLM output JSON.",
    "Volume target: 5,000-20,000 chunk→extraction pairs across diverse document types.",
    "Collection: Add a capture hook to FullExtractionPipeline.extract_unified() that logs (chunk_text, extraction_json) to backend/training_data/extraction_traces.jsonl when CAPTURE_TRAINING_DATA=true.",
    "Quality filter: Only keep traces that passed Pydantic schema validation and had a non-zero resolver merge rate.",
    "Synthetic augmentation: Use frontier model to paraphrase chunks (preserving entities) to 3-5x the dataset.",
    "Format: OpenAI chat format — system prompt = current UnifiedExtractionSignature docstring, user = chunk text, assistant = JSON extraction.",
])
add_subheading("Evaluation")
add_bullets([
    "Schema adherence: % of outputs parsing as valid ExtractionOutput (target: >99%)",
    "Entity recall: overlap of extracted entity labels vs. frontier model (target: >85%)",
    "Edge type accuracy: match rate on edge_type predictions (target: >80%)",
    "Latency: p50/p95 per-chunk extraction time (target: <5s p50)",
])
add_subheading("Files / Infra")
add_bullets([
    "backend/training/capture.py — extraction trace capture",
    "backend/training/finetune_extractor.py — LoRA fine-tuning script (Axolotl or Unsloth)",
    "backend/ingestion/dspy_programs.py — add SLM-backed provider option, toggled via settings.EXTRACTION_MODEL_PROVIDER",
])

# ── S2 ───────────────────────────────────────
add_section("S2", "Edge-Type Classifier SLM",
            "Specialized classifier (augments S1 or runs standalone)",
            "Medium", "High — fixes hallucinated edge types structurally")
add_subheading("Problem")
add_para(
    "The LLM picks suboptimal edge types from the 29-type vocabulary even with Literal constraints. Edges "
    "like AMALGAMATED_INTO get remapped to OTHER or RELATED_TO, losing specificity. Given (source_entity, "
    "context_sentence, target_entity), predict the best edge type — a classification problem."
)
add_subheading("Target SLM")
add_para(
    "A 3-4B parameter model (Phi-3-mini or Qwen2.5-3B) fine-tuned for single-turn classification. Input: "
    "sentence + source/target entity labels. Output: one of 22 semantic edge types (or OTHER with a "
    "suggested alternative). Runs at 100-200 predictions/sec on a single GPU."
)
add_subheading("Training Data")
add_bullets([
    "Source: Extraction traces from S1, decomposed into (context, source, target, edge_type) tuples.",
    "Volume target: 50,000+ labeled tuples.",
    "Negative sampling: 2-3 hard negatives per positive by picking confusable edge types (e.g., PART_OF vs INSTANCE_OF) via sentence embeddings.",
    "Human-in-the-loop: Review 500-1000 low-confidence cases, relabel, upweight during training.",
    "Format: 'Source: {src}\\nTarget: {tgt}\\nContext: {sentence}\\nEdge type:' → single-token answer.",
])
add_subheading("Evaluation")
add_bullets([
    "Top-1 accuracy (target: >85%)",
    "Top-3 accuracy with alternatives (target: >95%)",
    "Confusion matrix across edge types",
])
add_subheading("Files / Infra")
add_bullets([
    "backend/training/edge_classifier.py — training + inference",
    "backend/ingestion/extractor.py — optional post-processing stage that re-scores low-confidence edge types",
])

# ── S3 ───────────────────────────────────────
add_section("S3", "Extraction Quality Evaluator SLM",
            "Automated eval model (distilled LLM-as-judge)",
            "Medium", "High — unlocks continuous eval without frontier cost")
add_subheading("Problem")
add_para(
    "Measuring extraction quality requires manual review or calling a frontier model to grade each "
    "output (expensive). Without cheap automated eval, we can't catch regressions or compare variants at scale."
)
add_subheading("Target SLM")
add_para(
    "Llama-3.1-8B fine-tuned as a judge. Given (chunk_text, extraction_json), outputs structured scores "
    "for completeness, correctness, edge_type_quality, hallucination, overall — plus a list of specific issues."
)
add_subheading("Training Data")
add_bullets([
    "Source: (chunk, extraction, frontier_judge_score) triples with frontier model as gold judge.",
    "Volume target: 2,000-5,000 judge traces.",
    "Generation: Prompt GPT-5/Claude Opus with chunk + extraction + structured rubric; save judgment.",
    "Diversity: Perturb good extractions into bad ones (missing entities, wrong edges, hallucinated entities) so the judge learns to discriminate.",
    "Human calibration: Manually grade 100 extractions; refine rubric until frontier judge agrees.",
    "Format: Chat — system = rubric, user = chunk + extraction, assistant = JSON score.",
])
add_subheading("Evaluation")
add_bullets([
    "Correlation with human scores on a 50-example held-out set (target: Spearman >0.80)",
    "Agreement with frontier judge on held-out traces (target: >90% within ±0.1 on overall)",
])
add_subheading("Files / Infra")
add_bullets([
    "backend/eval/extraction_judge.py — inference + batch scoring",
    "backend/eval/eval_harness.py — aggregate metrics across the training set",
    "CI integration to catch extraction regressions per commit",
])

# ── S4 ───────────────────────────────────────
add_section("S4", "Cypher Query Generator SLM",
            "Text-to-Cypher code model (augments Q9)",
            "High", "Medium — reliable NL→Cypher without frontier cost")
add_subheading("Problem")
add_para(
    "trident_cypher works only when the frontier model writes correct Cypher. The LLM gets relationship "
    "direction wrong, misuses exists() (removed in Neo4j 5), and forgets provider_id scoping."
)
add_subheading("Target SLM")
add_para(
    "Qwen2.5-Coder-7B fine-tuned on Trident-schema-aware NL→Cypher pairs. Input: NL question + live graph "
    "schema. Output: parameterized Cypher with correct provider_id scoping and Neo4j 5 syntax."
)
add_subheading("Training Data")
add_bullets([
    "Source: Synthetic NL→Cypher pairs generated from the actual Trident schema.",
    "Volume target: 10,000-30,000 question+Cypher pairs.",
    "Step 1: Take the live graph schema for 3-5 providers.",
    "Step 2: Frontier model + few-shot examples generate 20-50 Cypher queries per schema (exact match, multi-hop, shortest path, aggregation, property filters, conditional joins).",
    "Step 3: Execute each against Neo4j; if it runs, keep it; if it errors, have the frontier model fix it.",
    "Step 4: Generate an NL question for each valid Cypher.",
    "Step 5: Augment with paraphrase variations.",
    "Neo4j 5 constraint: Filter out deprecated syntax (exists() on properties) via real-Neo4j validation.",
    "Format: Prompt injects live schema, followed by the question; assistant returns the Cypher.",
])
add_subheading("Evaluation")
add_bullets([
    "Execution rate: % of generated queries running without error (target: >95%)",
    "Answer match rate: % returning same rows as gold Cypher (target: >80%)",
    "Benchmarks: Spider, WikiSQL-adapted, custom Trident-schema test set (200 hand-written pairs)",
])
add_subheading("Files / Infra")
add_bullets([
    "backend/training/cypher_generator.py — training + inference",
    "backend/agent/tools.py — new trident_cypher_slm tool with frontier fallback",
    "backend/training/cypher_synth.py — synthetic data generation pipeline",
])

# ── S5 ───────────────────────────────────────
add_section("S5", "Semantic Resolver / Dedup Judge SLM",
            "Pairwise classification model (augments R7)",
            "Low", "Medium — faster and more accurate entity resolution")
add_subheading("Problem")
add_para(
    "Resolution compares candidate embeddings against the GN index with a fixed cosine threshold (0.85 "
    "for entities, 0.82 for concepts). Borderline cases (0.78-0.88) are all-merged or all-created, losing "
    "precision. The frontier-model fallback in R7 is expensive."
)
add_subheading("Target SLM")
add_para(
    "Phi-3-mini (3B) fine-tuned as a pairwise classifier: given two candidate entity descriptions, "
    "predict 'same' / 'different' / 'unclear'. Used as a lightweight tiebreaker for borderline matches."
)
add_subheading("Training Data")
add_bullets([
    "Source: Entity pairs from the graph that the current resolver merged, plus hard-negative pairs.",
    "Volume target: 20,000-50,000 pairs.",
    "Step 1: Sample all merged entity pairs from past ingestions (true positives).",
    "Step 2: For each merged pair, sample an entity with the same label prefix but different full label as a hard negative (e.g., 'Meridian MCP' vs 'Meridian MIG').",
    "Step 3: Frontier model generates description paraphrases — two paraphrases of the same entity = positive pair.",
    "Step 4: Include cross-domain negatives (entities from different providers sharing a label).",
    "Format: 'Entity A: {label_a}: {desc_a}\\nEntity B: {label_b}: {desc_b}\\nAre these the same?' → yes/no/unclear.",
])
add_subheading("Evaluation")
add_bullets([
    "Pairwise accuracy on held-out pairs (target: >92%)",
    "Abstention rate on 'unclear' cases — should flag borderline pairs for human review",
    "End-to-end impact: entity merge rate improvement when used as tiebreaker",
])
add_subheading("Files / Infra")
add_bullets([
    "backend/training/dedup_judge.py — training + inference",
    "backend/ingestion/resolver.py — integrate SLM tiebreaker for scores in [0.78, 0.88]",
])


# ── Cross-Cutting Concerns (SLMs) ────────────
doc.add_page_break()
doc.add_heading("SLM — Cross-Cutting Concerns", level=2)

add_subheading("Training Infrastructure")
add_para("Before implementing any SLM above, stand up shared training infra:")
add_bullets([
    "Training data capture: A common hook/logger in backend/training/capture.py wrapping DSPy calls; writes (prompt, response, metadata) to JSONL.",
    "Fine-tuning toolkit: Standardize on Unsloth (fastest) or Axolotl (most flexible). All SLM configs live in backend/training/configs/.",
    "Model serving: vLLM or Ollama for local inference; SLMs exposed via a common interface so dspy_programs.py can swap providers.",
    "Versioning: Each model gets a semantic version, stored in HuggingFace Hub or local MinIO. settings.EXTRACTION_MODEL=slm-extractor:v1.2 selects which model to use.",
])

add_subheading("Quality Gating")
add_para("Every SLM deployment must pass a canary stage:")
add_bullets([
    "Shadow-mode: run SLM alongside frontier for 7 days, log disagreements.",
    "A/B test: 10% of traffic on SLM, compare eval metrics.",
    "Full rollout: only if SLM quality is within 5% of frontier on the eval set.",
])

add_subheading("Cost Model")
add_matrix(
    ["Task", "Frontier cost/call", "SLM cost/call (local)", "Speedup"],
    [
        ("Entity extraction (S1)", "~$0.02", "~$0.0005", "40x"),
        ("Edge classification (S2)", "~$0.005", "~$0.0001", "50x"),
        ("Quality eval (S3)", "~$0.015", "~$0.0003", "50x"),
        ("Cypher generation (S4)", "~$0.008", "~$0.0002", "40x"),
        ("Dedup judgment (S5)", "~$0.003", "~$0.0001", "30x"),
    ],
)
p = doc.add_paragraph()
p.add_run(
    "Assuming 1M extractions/month at current scale: ~$20K/month on frontier → ~$500/month on SLMs "
    "(+ one-time fine-tuning cost of $200-500 per model)."
).italic = True


# ── Implementation Notes ─────────────────────
doc.add_page_break()
doc.add_heading("Implementation Notes", level=1)

add_subheading("Recommended Order — Representation")
add_para("R1 (fix edges) → R2 (reranking) → R3 (temporal) → R5 (confidence) → R6 (weights)")
add_para("    ↓")
add_para("R4 (communities) → R9 (DRIFT search)")
add_para("    ↓")
add_para("R7 (adaptive resolution) → R8 (drift detection)")

add_subheading("Recommended Order — Querying")
add_para("Q1 (planner) → Q3 (hybrid search) → Q4 (provenance)")
add_para("    ↓")
add_para("Q2 (find_path) + Q5 (aggregate) + Q6 (compare)  ← quick wins")
add_para("    ↓")
add_para("Q7 (entity memory) → Q8 (impact) → Q9 (text2cypher) → Q10 (suggest)")

add_subheading("Recommended Order — SLMs")
add_para("Training infra (capture + serving)")
add_para("    ↓")
add_para("S3 (quality evaluator) — needed to validate other SLMs")
add_para("    ↓")
add_para("S1 (entity extractor) — highest ROI")
add_para("    ↓")
add_para("S2 (edge classifier) + S5 (dedup judge)  ← refinements")
add_para("    ↓")
add_para("S4 (cypher generator) — depends on Q9")

add_subheading("Testing Strategy")
add_para("Each improvement should be validated by:")
add_bullets([
    "Unit tests on the new/modified methods.",
    "A 'before/after' comparison query set (5-10 questions the current system answers poorly).",
    "Graph statistics comparison (node count, edge count, edge-to-node ratio).",
])

add_subheading("Monitoring")
add_para("Track these metrics per provider:")
add_bullets([
    "Entity merge rate (% of entities that resolved to existing nodes)",
    "Average resolution score (should be bimodal: clear matches vs clear non-matches)",
    "Query anchor hit rate (% of queries that find at least one anchor node)",
    "Graph-boost reranking delta (how much chunk ordering changes after reranking)",
    "Community count and average community size",
    "Agent tool call distribution (which tools are used most/least)",
    "Path query success rate (for Q2)",
    "Provenance coverage (% of answer claims with source evidence)",
])


# ── Save ─────────────────────────────────────
out_path = Path(__file__).parent / "trident-roadmap.docx"
doc.save(str(out_path))
print(f"Generated: {out_path}")
